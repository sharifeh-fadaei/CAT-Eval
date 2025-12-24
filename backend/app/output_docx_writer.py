from __future__ import annotations

from pathlib import Path
from typing import Dict, Any, List

from docx import Document


def _render_table(document: Document, evaluation_table: List[Dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Factor"
    hdr_cells[1].text = "Weight"
    hdr_cells[2].text = "Score"
    hdr_cells[3].text = "Severity"
    hdr_cells[4].text = "Issue"
    for row in evaluation_table:
        cells = table.add_row().cells
        cells[0].text = str(row.get("factor", ""))
        cells[1].text = str(row.get("weight", ""))
        cells[2].text = str(row.get("score_1_to_10", ""))
        cells[3].text = str(row.get("error_severity", ""))
        cells[4].text = str(row.get("whats_wrong_missing_invented", ""))


def write_single_docx(run_dir: Path, run_id: str, record: Dict[str, Any]) -> Path:
    run_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("HQ-CAT Single Evaluation", level=1)
    doc.add_paragraph(f"Sample: {record['sample_name']}")
    doc.add_paragraph(f"Model: {record['model']}")
    doc.add_paragraph(f"Final Score: {record['final_score']:.2f}")
    _render_table(doc, record["evaluation_table"])
    doc.add_heading("Improvement Report", level=2)
    for bullet in record.get("improvement_report", []):
        doc.add_paragraph(bullet, style="List Bullet")
    output_path = run_dir / f"single_{run_id}.docx"
    doc.save(output_path)
    return output_path


def write_batch_docx(run_dir: Path, run_id: str, model_key: str, records: List[Dict[str, Any]]) -> Path:
    run_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(f"HQ-CAT Batch Evaluation – {model_key}", level=1)
    for record in records:
        doc.add_heading(record.get("sample_name", "sample"), level=2)
        doc.add_paragraph(f"Model: {record.get('model')}")
        doc.add_paragraph(f"Final Score: {record.get('final_score'):.2f}")
        _render_table(doc, record.get("evaluation_table", []))
        doc.add_heading("Improvement Report", level=3)
        for bullet in record.get("improvement_report", []):
            doc.add_paragraph(bullet, style="List Bullet")
    output_path = run_dir / f"batch_{run_id}_{model_key}.docx"
    doc.save(output_path)
    return output_path


def write_combined_docx(run_dir: Path, run_id: str, model_records: Dict[str, List[Dict[str, Any]]]) -> Path:
    run_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("HQ-CAT Batch Evaluation – Combined", level=1)
    for model_key, records in model_records.items():
        doc.add_heading(f"Model: {model_key}", level=2)
        for record in records:
            doc.add_heading(record.get("sample_name", "sample"), level=3)
            doc.add_paragraph(f"Final Score: {record.get('final_score'):.2f}")
            _render_table(doc, record.get("evaluation_table", []))
            doc.add_heading("Improvement Report", level=4)
            for bullet in record.get("improvement_report", []):
                doc.add_paragraph(bullet, style="List Bullet")
    output_path = run_dir / f"batch_{run_id}_combined.docx"
    doc.save(output_path)
    return output_path
